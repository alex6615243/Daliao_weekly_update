from __future__ import annotations

import argparse
import copy
import json
import re
import shutil
import sys
import urllib.request
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from urllib.parse import urlparse
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from lxml import html


APP_DIR = Path(__file__).resolve().parent
CONFIG_PATH = APP_DIR / "config.json"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
URL_RE = re.compile(r"https?://[^\s，。；、）)]+")
DATE_IN_NAME_RE = re.compile(r"(?P<y>20\d{2})[._-](?P<m>\d{1,2})[._-](?P<d>\d{1,2})")
CN_NUM = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}
BOOK_NAMES = {"路": "路加福音", "太": "馬太福音", "可": "馬可福音", "約": "約翰福音", "徒": "使徒行傳"}


@dataclass
class PageData:
    web_date: date
    web_issue: int | None
    series: str
    morning_week: int
    message_title: str
    verse: str
    burden_title: str
    burden_body: str
    bible_start: str
    bible_end: str
    traffic_items: list[str]


def clean(text: str) -> str:
    text = text.replace("\xa0", " ").replace("＊", "*")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    return text.strip()


def next_sunday(today: date) -> date:
    return today + timedelta(days=(6 - today.weekday()) % 7)


def parse_target(value: str | None) -> date:
    return datetime.strptime(value, "%Y-%m-%d").date() if value else next_sunday(date.today())


def download(url: str, cache_path: Path | None = None) -> bytes:
    if cache_path:
        return cache_path.read_bytes()
    request = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 WeeklyBulletinUpdater/1.0"})
    with urllib.request.urlopen(request, timeout=30) as response:
        return response.read()


def visible_lines(raw: bytes) -> list[str]:
    root = html.fromstring(raw)
    for node in root.xpath("//script|//style|//noscript"):
        node.drop_tree()
    lines = [clean(x) for x in root.text_content().splitlines()]
    return [x for x in lines if x]


def find_index(lines: list[str], needle: str, start: int = 0) -> int:
    for i in range(start, len(lines)):
        if needle in lines[i]:
            return i
    raise ValueError(f"網頁缺少必要標題：{needle}")


def parse_web_date(lines: list[str]) -> tuple[date, int | None]:
    for line in lines[:80]:
        if "召會週訊第" not in line:
            continue
        compact = line.replace(" ", "")
        dm = re.search(r"(20\d{2})/(\d{1,2})/(\d{1,2})", compact)
        im = re.search(r"第([\d,]+)期", compact)
        if dm:
            return date(int(dm.group(1)), int(dm.group(2)), int(dm.group(3))), int(im.group(1).replace(",", "")) if im else None
    raise ValueError("找不到網頁週訊日期，已停止以避免做錯週次。")


def merge_url_lines(lines: list[str]) -> list[str]:
    merged: list[str] = []
    for line in lines:
        if re.fullmatch(r"https?://\S+", line) and merged:
            merged[-1] = merged[-1].rstrip("：: ") + "：" + line
        elif merged and (merged[-1].endswith("：") or merged[-1].endswith(":")):
            merged[-1] += line
        else:
            merged.append(line)
    return merged


def parse_traffic(lines: list[str]) -> list[str]:
    start = find_index(lines, "召會生活動態交通：") + 1
    end = find_index(lines, "每日靈糧", start)
    src = merge_url_lines(lines[start:end])
    output: list[str] = []
    category = ""
    in_south_b = False
    for line in src:
        if line.startswith("★"):
            category = line.lstrip("★").rstrip("：:")
            in_south_b = False
            continue
        if category == "照顧區":
            if re.match(r"(?:\d+[.、]\s*)?南B區[：:]", line, re.I):
                in_south_b = True
                remainder = re.sub(r"^(?:\d+[.、]\s*)?南B區[：:]\s*", "", line, flags=re.I)
                if remainder:
                    output.append(remainder)
                continue
            if re.match(r"\d+[.、]\s*[^：:]+[：:]", line) or re.match(r"南[ACD]區[：:]", line, re.I):
                in_south_b = False
                continue
            if not in_south_b:
                continue
        line = re.sub(r"^(?:\d+[.、]|[①②③④⑤⑥⑦⑧⑨⑩])\s*", "", line)
        line = line.lstrip("★").strip()
        if line and line not in {"全召會：", "兒童：", "中學：", "大學：", "書房：", "影音文字組："}:
            output.append(line)
    return [x for x in output if len(x) > 3]


def parse_reading_progress(lines: list[str]) -> tuple[str, int, str, str]:
    daily = find_index(lines, "每日靈糧")
    prayer = find_index(lines, "禱", daily)
    header_lines = lines[daily:prayer]
    header = next((x for x in header_lines if "晨興聖言" in x and "週" in x), "")
    match = re.search(r"[「『](.+?)[」』]晨興聖言第([一二三四五六七八九十\d]+)週", header)
    if not match:
        raise ValueError("無法辨識晨興聖言週次。")
    series = match.group(1).replace("（", "(").replace("）", ")")
    week_text = match.group(2)
    week = int(week_text) if week_text.isdigit() else CN_NUM.get(week_text, 0)
    progress = [re.sub(r"^讀經進度[：:]\s*", "", x) for x in lines[prayer:] if x.startswith("讀經進度")]
    progress = [x for x in progress if "複習" not in x]
    if len(progress) < 2:
        raise ValueError("讀經進度不足，無法計算起訖。")
    return series, week, progress[0].split("~", 1)[0], progress[-1].split("~", 1)[-1]


def parse_outline(lines: list[str]) -> tuple[str, str, str, str]:
    start = find_index(lines, "綱要") + 1
    end = find_index(lines, "召會生活交通", start)
    block = lines[start:end]
    message_pos = next(i for i, x in enumerate(block) if re.match(r"第[一二三四五六七八九十\d]+篇", x))
    message = block[message_pos]
    verse = block[message_pos + 1]
    candidates: list[tuple[str, str]] = []
    for i in range(message_pos + 2, len(block) - 1):
        if 4 <= len(block[i]) <= 34 and len(block[i + 1]) >= 100 and not re.match(r"[壹貳參肆伍陸柒捌玖拾一二三四五六七八九十]、?", block[i]):
            candidates.append((block[i], block[i + 1]))
    if not candidates:
        raise ValueError("找不到負擔的話標題與內文。")
    burden_title, burden_body = candidates[-1]
    return message, verse, burden_title, burden_body


def parse_page(raw: bytes) -> PageData:
    lines = visible_lines(raw)
    web_date, web_issue = parse_web_date(lines)
    series, week, bible_start, bible_end = parse_reading_progress(lines)
    message, verse, burden_title, burden_body = parse_outline(lines)
    return PageData(web_date, web_issue, series, week, message, verse, burden_title, burden_body, bible_start, bible_end, parse_traffic(lines))


def filename_date(path: Path) -> date | None:
    match = DATE_IN_NAME_RE.search(path.name)
    if not match:
        return None
    try:
        return date(int(match.group("y")), int(match.group("m")), int(match.group("d")))
    except ValueError:
        return None


def choose_template(folder: Path, approved: str, target: date, explicit: str | None) -> Path:
    if explicit:
        result = Path(explicit).resolve()
        if not result.exists():
            raise FileNotFoundError(result)
        return result
    candidates = []
    for path in folder.glob("*.docx"):
        d = filename_date(path)
        if d and d <= target and not path.name.startswith("~$"):
            candidates.append((d, path.stat().st_mtime, path))
    if candidates:
        return max(candidates, key=lambda x: (x[0], x[1]))[2]
    fallback = folder / approved
    if fallback.exists():
        return fallback
    raise FileNotFoundError("找不到可用的 Word 範本。")


def unique_output(folder: Path, pattern: str, target: date) -> Path:
    base = folder / pattern.format(date=target.strftime("%Y.%m.%d"))
    if not base.exists():
        return base
    for i in range(1, 100):
        alt = folder / f"{base.stem}_自動更新_{i}{base.suffix}"
        if not alt.exists():
            return alt
    raise FileExistsError("同日期輸出檔過多，請先整理檔名。")


def replace_paragraph_text(paragraph, text: str) -> None:
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def split_balanced(text: str, max_first: int = 32) -> tuple[str, str]:
    if len(text) <= max_first:
        return text, ""
    cut = max(text.rfind("，", 0, max_first + 1), text.rfind("；", 0, max_first + 1), text.rfind(" ", 0, max_first + 1))
    if cut < max_first // 2:
        cut = max_first
    else:
        cut += 1
    return text[:cut].rstrip(), text[cut:].lstrip()


def cn_chapter_to_arabic(text: str) -> tuple[str, int, int]:
    match = re.match(r"([\u4e00-\u9fff]+?)([一二三四五六七八九十百]+)(\d+)$", text)
    if not match:
        raise ValueError(f"無法辨識讀經範圍：{text}")
    book = BOOK_NAMES.get(match.group(1), match.group(1))
    chars = match.group(2)
    if chars == "十": number = 10
    elif "十" in chars:
        a, _, b = chars.partition("十")
        number = (CN_NUM.get(a, 1) * 10) + CN_NUM.get(b, 0)
    else: number = CN_NUM.get(chars, 0)
    return book, number, int(match.group(3))


def bible_summary(start: str, end: str) -> str:
    b1, c1, v1 = cn_chapter_to_arabic(start)
    b2, c2, v2 = cn_chapter_to_arabic(end)
    return f"聖經追求進度：{b1}{to_cn(c1)}章{v1}節至{b2}{to_cn(c2)}章{v2}節。"


def to_cn(number: int) -> str:
    rev = {v: k for k, v in CN_NUM.items() if v < 10}
    if number < 10: return rev[number]
    if number == 10: return "十"
    if number < 20: return "十" + rev.get(number % 10, "")
    return rev[number // 10] + "十" + rev.get(number % 10, "")


def latest_item_date(text: str, target: date) -> date | None:
    dates: list[date] = []
    for month, day in re.findall(r"(?<!\d)(\d{1,2})月(\d{1,2})日", text):
        try: dates.append(date(target.year, int(month), int(day)))
        except ValueError: pass
    for month, day in re.findall(r"(?<!\d)(\d{1,2})月(\d{1,2})(?:~|～|至|-)(\d{1,2})日", text):
        try: dates.append(date(target.year, int(month), int(day)))
        except ValueError: pass
    for month, _start, end in re.findall(r"(?<!\d)(\d{1,2})月(\d{1,2})(?:~|～|至|-)(\d{1,2})日", text):
        try: dates.append(date(target.year, int(month), int(end)))
        except ValueError: pass
    return max(dates) if dates else None


def item_key(text: str) -> str:
    head = re.split(r"[：:，,。]", text, maxsplit=1)[0]
    return re.sub(r"\s|20\d{2}年|\d+月\d+日", "", head)[:28]


def should_keep_old(text: str, target: date, keep_undated: bool) -> bool:
    d = latest_item_date(text, target)
    return keep_undated if d is None else d >= target


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def clear_content_keep_properties(p_element) -> None:
    for child in list(p_element):
        if child.tag != qn("w:pPr"):
            p_element.remove(child)


def append_text_and_links(document: Document, p_element, text: str) -> None:
    colon = text.find("：")
    if colon > 0:
        append_segment_and_links(document, p_element, text[:colon], bold=True)
        append_segment_and_links(document, p_element, text[colon:], bold=False)
    else:
        append_segment_and_links(document, p_element, text, bold=False)


def append_segment_and_links(document: Document, p_element, text: str, bold: bool) -> None:
    pos = 0
    for match in URL_RE.finditer(text):
        add_run_xml(p_element, text[pos:match.start()], bold=bold)
        url = match.group(0)
        rel_id = document.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel_id)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
        underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
        size = OxmlElement("w:sz"); size.set(qn("w:val"), "27")
        size_cs = OxmlElement("w:szCs"); size_cs.set(qn("w:val"), "27")
        rpr.extend([color, underline, size, size_cs])
        if bold:
            bold_node = OxmlElement("w:b")
            rpr.append(bold_node)
        run.append(rpr)
        t = OxmlElement("w:t"); t.text = url; run.append(t)
        hyperlink.append(run); p_element.append(hyperlink)
        pos = match.end()
    add_run_xml(p_element, text[pos:], bold=bold)


def add_run_xml(p_element, text: str, bold: bool = False) -> None:
    if not text: return
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    size = OxmlElement("w:sz"); size.set(qn("w:val"), "27")
    size_cs = OxmlElement("w:szCs"); size_cs.set(qn("w:val"), "27")
    rpr.extend([size, size_cs])
    if bold:
        bold_node = OxmlElement("w:b")
        rpr.append(bold_node)
    run.append(rpr)
    t = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace(): t.set(qn("xml:space"), "preserve")
    t.text = text; run.append(t); p_element.append(run)


def update_document(template: Path, output: Path, target: date, data: PageData, config: dict) -> None:
    shutil.copy2(template, output)
    doc = Document(output)
    source_paragraphs = doc.paragraphs
    prayer_heading = next(p for p in source_paragraphs if "《代禱事項》" in p.text)
    prayer_before = prayer_heading._element.xml

    issue_match = re.search(r"(\d+)期", source_paragraphs[0].text)
    source_date = filename_date(template)
    if source_date is None:
        header_date = re.search(r"主後(\d{4})年(\d{1,2})月(\d{1,2})日", source_paragraphs[1].text)
        if header_date:
            source_date = date(*(int(value) for value in header_date.groups()))
    source_date = source_date or target - timedelta(days=7)
    issue = int(issue_match.group(1)) + max(0, (target - source_date).days // 7) if issue_match else data.web_issue
    replace_paragraph_text(source_paragraphs[0], f"   {issue}期")
    replace_paragraph_text(source_paragraphs[1], re.sub(r"主後\d{4}年\d{1,2}月\d{1,2}日", f"主後{target.year}年{target.month:02d}月{target.day:02d}日", source_paragraphs[1].text))
    replace_paragraph_text(source_paragraphs[2], f"《{data.series}》")
    replace_paragraph_text(source_paragraphs[3], data.message_title)
    verse1, verse2 = split_balanced(data.verse, 34)
    replace_paragraph_text(source_paragraphs[4], verse1)
    replace_paragraph_text(source_paragraphs[5], verse2)
    title1, title2 = split_balanced(data.burden_title, 10)
    replace_paragraph_text(source_paragraphs[7], "    " + title1)
    replace_paragraph_text(source_paragraphs[8], title2)
    replace_paragraph_text(source_paragraphs[9], "        " + data.burden_body)

    paragraphs = doc.paragraphs
    report_heading = next(p for p in paragraphs if "《本週報告事項》" in p.text)
    prayer_heading = next(p for p in paragraphs if "《代禱事項》" in p.text)
    start = paragraphs.index(report_heading) + 1
    end = paragraphs.index(prayer_heading)
    report_paras = [p for p in paragraphs[start:end] if p.text.strip() and p.text.strip() != "。"]
    if len(report_paras) < 3:
        raise ValueError("範本中的報告事項結構不完整。")
    morning = f"全召會晨興追求進度: 半年度訓練「{data.series}」第{to_cn(data.morning_week)}篇。"
    replace_paragraph_text(report_paras[0], morning)
    replace_paragraph_text(report_paras[1], bible_summary(data.bible_start, data.bible_end))
    replace_paragraph_text(report_paras[2], config["prayer_meeting_text"])

    fixed_keys = {item_key(morning), item_key(report_paras[1].text), item_key(config["prayer_meeting_text"])}
    old_keep = []
    for p in report_paras[3:]:
        if should_keep_old(p.text, target, config.get("keep_undated_old_items", True)):
            old_keep.append(p)
        else:
            remove_paragraph(p)
    old_by_key = {item_key(p.text): p for p in old_keep}
    for new_text in data.traffic_items:
        key = item_key(new_text)
        if key in fixed_keys:
            continue
        if key in old_by_key:
            p = old_by_key.pop(key)
            clear_content_keep_properties(p._element)
            append_text_and_links(doc, p._element, new_text)
        else:
            sample = report_paras[2]._element
            new_p = copy.deepcopy(sample)
            clear_content_keep_properties(new_p)
            append_text_and_links(doc, new_p, new_text)
            prayer_heading._element.addprevious(new_p)

    # Normalize every report item after insertion. This also corrects retained
    # paragraphs whose direct formatting came from an older bulletin.
    paragraphs = doc.paragraphs
    report_heading = next(p for p in paragraphs if "《本週報告事項》" in p.text)
    prayer_heading = next(p for p in paragraphs if "《代禱事項》" in p.text)
    for paragraph in paragraphs[paragraphs.index(report_heading) + 1:paragraphs.index(prayer_heading)]:
        if not paragraph.text.strip() or paragraph.text.strip() == "。":
            continue
        text = paragraph.text
        clear_content_keep_properties(paragraph._element)
        append_text_and_links(doc, paragraph._element, text)

    if prayer_heading._element.xml != prayer_before:
        raise RuntimeError("《代禱事項》標題發生非預期變更。")
    doc.save(output)


def validate_docx(path: Path) -> None:
    with ZipFile(path) as package:
        if package.testzip() is not None:
            raise RuntimeError("輸出 Word 檔案結構損壞。")
        xml = package.read("word/document.xml").decode("utf-8")
        if "《代禱事項》" not in xml or "《本週報告事項》" not in xml:
            raise RuntimeError("輸出缺少必要段落。")


def write_preview(path: Path, template: Path, output: Path, target: date, data: PageData) -> None:
    lines = [
        f"目標日期：{target}", f"網站日期：{data.web_date}", f"使用範本：{template}", f"輸出檔案：{output}",
        f"主題：《{data.series}》", f"信息：{data.message_title}", f"負擔：{data.burden_title}",
        f"讀經：{data.bible_start} 至 {data.bible_end}", "", "網頁報告事項：",
    ] + [f"- {x}" for x in data.traffic_items]
    path.write_text("\n".join(lines), encoding="utf-8-sig")


def main() -> int:
    parser = argparse.ArgumentParser(description="大寮召會週訊一鍵更新")
    parser.add_argument("--target", help="目標主日，格式 YYYY-MM-DD；省略時自動取本週日")
    parser.add_argument("--template", help="指定 Word 範本")
    parser.add_argument("--html", help="使用本機 HTML 測試，不連線")
    parser.add_argument("--force", action="store_true", help="允許網站日期與目標主日不一致（僅供測試）")
    parser.add_argument("--output-dir", help="另指定輸出資料夾（測試用）")
    args = parser.parse_args()
    config = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    target = parse_target(args.target)
    folder = Path(args.output_dir) if args.output_dir else Path(config["bulletin_folder"])
    folder.mkdir(parents=True, exist_ok=True)
    raw = download(config["website"], Path(args.html) if args.html else None)
    data = parse_page(raw)
    if data.web_date != target and not args.force:
        raise RuntimeError(f"網站目前是 {data.web_date:%Y/%m/%d}，目標是 {target:%Y/%m/%d}；已安全停止，請等網站更新後再執行。")
    template_folder = Path(config["bulletin_folder"])
    template = choose_template(template_folder, config["approved_template"], target, args.template)
    output = unique_output(folder, config["output_name"], target)
    preview = folder / f"{output.stem}-更新摘要.txt"
    update_document(template, output, target, data, config)
    validate_docx(output)
    write_preview(preview, template, output, target, data)
    print(f"完成：{output}")
    print(f"摘要：{preview}")
    print("原始檔未被修改。請開啟 Word 做最後人工確認。")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"\n更新失敗：{exc}", file=sys.stderr)
        print("未覆寫任何原始檔。", file=sys.stderr)
        raise SystemExit(1)
